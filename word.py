import streamlit as st
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image, ImageDraw, ImageFont, ImageOps
from PIL.ExifTags import TAGS
import pillow_heif
import io
import os
import re
from datetime import datetime

# 註冊 HEIC 格式支援
pillow_heif.register_heif_opener()

# 初始化 session_state
if "uploaded_photos_list" not in st.session_state:
    st.session_state.uploaded_photos_list = []
    
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0

# -----------------------------------------
# 照片處理與日期讀取
# -----------------------------------------
def get_photo_date(image_bytes, file_name):
    """ 從照片的 EXIF 資訊或檔名讀取拍攝日期 """
    try:
        image = Image.open(io.BytesIO(image_bytes))
        exif = image._getexif()
        if exif:
            for tag, value in exif.items():
                decoded = TAGS.get(tag, tag)
                if decoded in ("DateTimeOriginal", "DateTime"):
                    date_str = value.split()[0].replace(":", "/")
                    return date_str, True
    except Exception:
        pass
    
    match = re.search(r"(\d{4})[-_ /]?(\d{2})[-_ /]?(\d{2})", file_name)
    if match:
        try:
            date_part = f"{match.group(1)}/{match.group(2)}/{match.group(3)}"
            dt = datetime.strptime(date_part, "%Y/%m/%d")
            return dt.strftime("%Y/%m/%d"), False
        except Exception:
            pass
            
    return datetime.today().strftime("%Y/%m/%d"), False

def resize_and_compress_image(image_bytes, date_str, is_from_exif, print_watermark):
    """
    處理照片尺寸與浮水印：
    - 自動修正手機拍攝時的 EXIF 旋轉問題。
    - 橫式照片：置中裁切填滿。
    - 直式照片：等比例縮放且限制高度不超出儲存格。
    - 當 print_watermark 為 True 且日期不為空時，才會在右下角繪製時間文字。
    """
    raw_img = Image.open(io.BytesIO(image_bytes))
    img = ImageOps.exif_transpose(raw_img)
    
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
        
    target_pixel_w = 945  # 8.0cm 對應 pixel (300 DPI)
    target_ratio = 8.0 / 6.15
    target_pixel_h = int(target_pixel_w / target_ratio)  # 約 726 像素 (6.15cm 安全高度)
    
    img_w, img_h = img.size
    
        # 縮放控制：直式照片限制最高為表格安全高度，絕對不超出格線
    if img_w < img_h:
        ratio_w = target_pixel_w / img_w
        ratio_h = target_pixel_h / img_h
        scale_ratio = min(ratio_w, ratio_h)
        new_w = int(img_w * scale_ratio)
        new_h = int(img_h * scale_ratio)
        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
    else:
        # 橫式照片：置中裁切並縮放
        current_ratio = img_w / img_h
        if current_ratio > target_ratio:
            crop_w = int(target_ratio * img_h)
            offset = (img_w - crop_w) // 2
            img = img.crop((offset, 0, img_w - offset, img_h))
        elif current_ratio < target_ratio:
            crop_h = int(img_w / target_ratio)
            offset = (img_h - crop_h) // 2
            img = img.crop((0, offset, img_w, img_h - offset))
        img = img.resize((target_pixel_w, target_pixel_h), Image.Resampling.LANCZOS)

    # 使用者手動決定是否蓋上時間浮水印
    if print_watermark and date_str.strip() != "":
        draw = ImageDraw.Draw(img)
        
        dpi = img.info.get('dpi', (300, 300))[0]
        if dpi <= 0:
            dpi = 300
        font_size = max(40, int(0.35 * dpi))
        border_thickness = max(4, int(font_size * 0.08))
        
        try:
            font = ImageFont.truetype("msjh.ttc", font_size)  
        except IOError:
            try:
                font = ImageFont.truetype("Arial.ttf", font_size)
            except IOError:
                font = ImageFont.load_default()

        try:
            text_bbox = draw.textbbox((0, 0), date_str, font=font)
            text_w = text_bbox[2] - text_bbox[0]
            text_h = text_bbox[3] - text_bbox[1]
        except AttributeError:
            text_w, text_h = draw.textsize(date_str, font=font)

        w, h = img.size
        margin = int(0.15 * dpi)
        x = w - text_w - margin
        y = h - text_h - margin
        
        border_thickness = 3
        for dx in range(-border_thickness, border_thickness + 1):
            for dy in range(-border_thickness, border_thickness + 1):
                if dx != 0 or dy != 0:
                    draw.text((x + dx, y + dy), date_str, font=font, fill="black")
                    
        draw.text((x, y), date_str, font=font, fill="white")

    out_io = io.BytesIO()
    img.save(out_io, format="JPEG", quality=95, dpi=(300, 300))
    out_io.seek(0)
    return out_io

# -----------------------------------------
# 固定欄寬與不拆頁 (Word 工具函式)
# -----------------------------------------
def set_cell_width(cell, width_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_row_height(row, height_cm):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

# -----------------------------------------
# 建立 Word 文件
# -----------------------------------------
def create_report(grouped_photos):
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    
    style = doc.styles['Normal']
    style.font.name = '標楷體'
    style.element.rPr.get_or_add_rFonts().set(qn('w:hint'), 'eastAsia')
    style.font.size = Pt(12)

    sorted_dates = sorted(grouped_photos.keys())
    
    for date_idx, date_str in enumerate(sorted_dates):
        photos = grouped_photos[date_str]
        total_photos = len(photos)
        
        for page_idx in range(0, max(1, total_photos), 6):
            if date_idx > 0 or page_idx > 0:
                doc.add_page_break()
                
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_before = Pt(0)
            title.paragraph_format.space_after = Pt(6)
            run = title.add_run(f"監造報表 ({date_str})")
            run.bold = True
            run.font.size = Pt(16)
            
            table = doc.add_table(rows=6, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            page_photos = photos[page_idx:page_idx + 6]
            
            for block_idx in range(3):
                r_big = block_idx * 2
                r_small = block_idx * 2 + 1
                
                row_big = table.rows[r_big]
                row_small = table.rows[r_small]
                
                set_row_height(row_big, 6.6)
                set_row_height(row_small, 0.8)
                set_row_cant_split(row_big)
                set_row_cant_split(row_small)
                
                for c_idx in range(2):
                    cell_big = row_big.cells[c_idx]
                    cell_small = row_small.cells[c_idx]
                    
                    set_cell_width(cell_big, 8.0)
                    set_cell_width(cell_small, 8.0)
                    
                    cell_big.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell_small.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    photo_pos = block_idx * 2 + c_idx
                    
                    p_para = cell_big.paragraphs[0]
                    p_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_para.paragraph_format.space_before = Pt(0)
                    p_para.paragraph_format.space_after = Pt(0)
                    
                    t_para = cell_small.paragraphs[0]
                    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    t_para.paragraph_format.space_before = Pt(0)
                    t_para.paragraph_format.space_after = Pt(0)
                    
                    if photo_pos < len(page_photos):
                        photo_data = page_photos[photo_pos]
                        
                        # 根據勾選狀態以及輸入日期，決定是否在圖片上加上浮水印
                        processed_img_io = resize_and_compress_image(
                            photo_data["raw_bytes"] or photo_data.get("bytes"), 
                            photo_data["date"], 
                            False,
                            photo_data["print_watermark"]
                        )
                        p_run = p_para.add_run()
                        
                        # 判斷直橫式： 高度限制縮小放入儲存格；橫式則固定
                        temp_img = Image.open(processed_img_io)
                        w, h = temp_img.size
                        if w < h:
                            # 直式照片等比例縮小，高度設為 6.15 cm 
                            p_run.add_picture(processed_img_io, height=Cm(6.15))
                        else:
                            # 橫式照片固定 8 cm 
                            p_run.add_picture(processed_img_io, width=Cm(8.0))
                        
                        t_run = t_para.add_run(photo_data["description"])
                        t_run.font.size = Pt(10)
                        
    out_doc_io = io.BytesIO()
    doc.save(out_doc_io)
    out_doc_io.seek(0)
    return out_doc_io


# -----------------------------------------
# Streamlit 前端網頁介面
# -----------------------------------------
st.title("偷懶小幫手 ฅ^•ﻌ•^ฅ")
st.caption("⚠️ 提醒：使用時間相機app是最好的喔！雖然還是可以幫你打印時間")

if "uploaded_photos_list" not in st.session_state:
    st.session_state.uploaded_photos_list = []

# 初始化控制清空 key
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0

uploaded_files = st.file_uploader(
    "上傳照片", 
    type=["jpg", "jpeg", "png", "heic"], 
    accept_multiple_files=True,
    key=f"file_uploader_{st.session_state.uploader_key}"
)

if uploaded_files:
    existing_names = {p["name"] for p in st.session_state.uploaded_photos_list}
    for f in uploaded_files:
        if f.name not in existing_names:
            file_bytes = f.read()
            
            # 嘗試抓取拍攝日期
            date_str, is_from_exif = get_photo_date(file_bytes, f.name)
            
            # 如果不是從真實 EXIF 抓到的，或者抓到的是今天(錯誤)日期，直接判定為「未偵測到日期」
            if not is_from_exif:
                # 預設不壓浮水印
                processed_io = resize_and_compress_image(file_bytes, date_str="", is_from_exif=False, print_watermark=False)
                default_display_date = "" 
            else:
                # 有抓到正確 EXIF，一樣預設不壓浮水印
                # 
                processed_io = resize_and_compress_image(file_bytes, date_str, is_from_exif, print_watermark=False)
                default_display_date = date_str

            # 將初始值儲存在 session_state 中
            st.session_state.uploaded_photos_list.append({
                "name": f.name,
                "bytes": processed_io.getvalue(),
                "display_date": default_display_date,
                "print_watermark": False,  # 呈現的勾選框預設為關閉
                "description": ""
            })

# 圖片管理與編輯介面
if st.session_state.uploaded_photos_list:
    
    # 一鍵刪除功能鍵清除日期與文字
    col_btn1, col_btn2 = st.columns([1, 3])
    with col_btn1:
        if st.button("🗑️ 一鍵刪除", type="secondary", use_container_width=True):
            for idx in range(len(st.session_state.uploaded_photos_list)):
                date_key = f"date_input_{idx}"
                desc_key = f"desc_input_{idx}"
                if date_key in st.session_state:
                    del st.session_state[date_key]
                if desc_key in st.session_state:
                    del st.session_state[desc_key]
            
            # 2. 清除照片資料與重置上傳鍵
            st.session_state.uploaded_photos_list = []
            st.session_state.uploader_key += 1
            st.rerun()

    # 防呆:點擊別處或按 Enter 時才存檔
    def update_photo_data(index, field_key, session_key):
        st.session_state.uploaded_photos_list[index][field_key] = st.session_state[session_key]

    # 分組資料
    grouped_ui = {}
    for idx, photo in enumerate(st.session_state.uploaded_photos_list):
        grouped_ui.setdefault(photo["display_date"], []).append((idx, photo))
        
    # 排序: 讓沒有時間的排在最上面
    sorted_groups = sorted(grouped_ui.keys(), key=lambda x: (x != "", x))
    
    for d_str in sorted_groups:
        
        # 沒有確定的時間時，直接不變動
        if d_str == "":
            st.markdown("請手動輸入時間（日期判定錯誤/未偵測到日期）")
            
            for original_idx, photo in grouped_ui[d_str]:
                col1, col2 = st.columns([1, 3])  
                with col1:
                    st.image(photo["bytes"], use_container_width=True)
                with col2:
                    date_key = f"date_input_{original_idx}"
                    st.text_input(
                        "補充拍攝日期_自己手動吧~", 
                        value=photo["display_date"], 
                        placeholder="輸入日期（例：YYYY/MM/DD or MM/DD）",
                        key=date_key,
                        on_change=update_photo_data,
                        args=(original_idx, "display_date", date_key)
                    )
                    
                    # 打印時間浮水印勾選框，維持既有頁面形式
                    wm_key = f"wm_input_{original_idx}"
                    st.checkbox(
                        "打印時間浮水印",
                        value=photo["print_watermark"],
                        key=wm_key,
                        on_change=update_photo_data,
                        args=(original_idx, "print_watermark", wm_key)
                    )
                    
                    desc_key = f"desc_input_{original_idx}"
                    st.text_input(
                        "照片說明", 
                        value=photo["description"], 
                        key=desc_key,
                        on_change=update_photo_data,
                        args=(original_idx, "description", desc_key)
                    )
                st.write("---")
                
        # 有確定的時間，就會有縮放功能
        else:
            photo_count = len(grouped_ui[d_str])
            # 使用 expander 建立可折疊區塊，預設設為 False
            with st.expander(f"📅 日期：{d_str} (共 {photo_count} 張照片)", expanded=False):
                
                for original_idx, photo in grouped_ui[d_str]:
                    col1, col2 = st.columns([1, 3])
                    with col1:
                        st.image(photo["bytes"], use_container_width=True)
                    with col2:
                        date_key = f"date_input_{original_idx}"
                        st.text_input(
                            "修改日期_如有問題自己改喔~", 
                            value=photo["display_date"], 
                            key=date_key,
                            on_change=update_photo_data,
                            args=(original_idx, "display_date", date_key)
                        )
                        
                        # 打印時間浮水印勾選框，維持既有頁面形式
                        wm_key = f"wm_input_{original_idx}"
                        st.checkbox(
                            "打印時間浮水印",
                            value=photo["print_watermark"],
                            key=wm_key,
                            on_change=update_photo_data,
                            args=(original_idx, "print_watermark", wm_key)
                        )
                        
                        desc_key = f"desc_input_{original_idx}"
                        st.text_input(
                            "照片說明", 
                            value=photo["description"], 
                            key=desc_key,
                            on_change=update_photo_data,
                            args=(original_idx, "description", desc_key)
                        )
                    st.write("---")

    # 封裝分組數據傳給報表生成函數
    final_grouped_photos = {}
    for photo in st.session_state.uploaded_photos_list:
        
        # 安全抓取圖片
        img_data = photo.get("raw_bytes") or photo.get("bytes") or photo.get("preview_bytes")
        
        report_photo_node = {
            "name": photo["name"],
            "raw_bytes": img_data,          
            "date": photo["display_date"], 
            "print_watermark": photo["print_watermark"],
            "description": photo["description"]
        }
        final_grouped_photos.setdefault(photo["display_date"], []).append(report_photo_node)

    try:
        word_file = create_report(final_grouped_photos)
        
        # 獲取今天的日期及時間字串，避免 import 衝突
        import datetime
        current_date_str = datetime.datetime.now().strftime('%Y%m%d')
        
        st.download_button(
            label="🚀 按下去後就可以節省很多時間喔",
            data=word_file,
            file_name=f"監造報表_{current_date_str}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
    except Exception as e:
        st.error(f"產生 Word 報表失敗，給我檢查清楚資料有沒有好嗎 :(：{e}")
